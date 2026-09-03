import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

PRIMARY_COLOR = RGBColor(30, 58, 138)     # #1E3A8A Dark Navy
SECONDARY_COLOR = RGBColor(37, 99, 235)   # #2563EB Blue
DARK_GRAY = RGBColor(51, 65, 85)          # #334155 Slate
LIGHT_BG_HEX = "F1F5F9"
CODE_BG_HEX = "0F172A"

def set_page_margins(doc):
    for s in doc.sections:
        s.top_margin = Inches(0.75)
        s.bottom_margin = Inches(0.75)
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)

def add_heading_styled(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Segoe UI"
    if level == 1:
        r.font.size = Pt(18)
        r.font.color.rgb = PRIMARY_COLOR
    elif level == 2:
        r.font.size = Pt(14)
        r.font.color.rgb = SECONDARY_COLOR
    elif level == 3:
        r.font.size = Pt(12)
        r.font.color.rgb = DARK_GRAY
    return p

def add_callout_box(doc, title, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(7.0)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{LIGHT_BG_HEX}"/>')
    tcPr.append(shd)
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="2563EB"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(borders)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r_title = p.add_run(f"📌 {title}\n")
    r_title.bold = True
    r_title.font.name = "Segoe UI"
    r_title.font.color.rgb = PRIMARY_COLOR
    r_text = p.add_run(text)
    r_text.font.name = "Segoe UI"
    r_text.font.size = Pt(10.5)

# ==============================================================================
# 1. BUILD FILE 1: ReconAI_Complete_Master_Prompt.docx
# ==============================================================================
print("Generating File 1: ReconAI_Complete_Master_Prompt.docx...")
doc_prompt = docx.Document()
set_page_margins(doc_prompt)

p1 = doc_prompt.add_paragraph()
r1 = p1.add_run("ReconAI — Complete Master Architecture & Development Prompt")
r1.bold = True
r1.font.name = "Segoe UI"
r1.font.size = Pt(22)
r1.font.color.rgb = PRIMARY_COLOR

p1_sub = doc_prompt.add_paragraph()
r1_sub = p1_sub.add_run("Full Autonomous Build Specification for Google Antigravity & AI Coding Agents")
r1_sub.font.name = "Segoe UI"
r1_sub.font.size = Pt(12)
r1_sub.font.italic = True
r1_sub.font.color.rgb = RGBColor(100, 116, 139)

add_callout_box(doc_prompt, "Document Purpose", "This document contains the complete, unabridged Master Prompt to generate or regenerate the entire ReconAI application from scratch. It includes all architecture blueprints, multi-pass reconciliation logic, universal document ingestion, Gemini 2.5 multimodal OCR, forensic audit rules, and packaging specs.")

master_prompt_text = """
# Master Prompt: ReconAI — Bank Reconciliation & Expense Audit Assistant

## 1. System Identity & High-Level Objective
You are a senior Principal Python Architect and Full-Stack Desktop Engineer. Build "ReconAI" — a production-grade, enterprise desktop Bank Reconciliation and Forensic Expense Audit Assistant engineered specifically for Chartered Accountants (CAs), CFOs, and Audit Firms.

ReconAI ingests multi-page bank statements (3,700+ transactions) and client accounting records (Tally Prime exports, Excel workbooks, CSV, Word .docx, scanned PDFs, and phone photos of handwritten vouchers), reconciles transactions using a deterministic and semantic multi-pass matching cascade, audits vouchers for forensic tax and regulatory risks (GSTIN, TDS, Section 40A(3), duplicate vouchers), and generates real-time Bank Reconciliation Statements (BRS) with formatted 7-Tab Excel workbooks and professional CA signed PDF deliverables.

---

## 2. Technology Stack & Key Dependencies
- Python: 3.11+ / 3.14+ 64-bit
- Desktop GUI: PyQt6 (QDarkTheme styling with Light/Dark toggle, Stepper Navigation, QStackedWidget)
- Tabular Data: pandas, openpyxl, xlsxwriter
- Text & Vector Matching: rapidfuzz (token_sort_ratio, token_set_ratio)
- PDF Extraction: pdfplumber, pypdfium2
- Document Ingestion: python-docx (for Word .docx tables)
- Multimodal Vision AI: google-genai SDK (Gemini 2.5 Flash for image/handwritten voucher OCR)
- Database Persistence: sqlite3 (session history, client master records, audit trails)
- Reporting Deliverables: reportlab (Executive PDF generation) & openpyxl (7-tab styled workbook)
- Test Suite: pytest (100% automated coverage)
- Binary Packaging: PyInstaller (.exe standalone bundle)

---

## 3. Core Architectural Modules

### Module 1: Universal Document Ingestion Subsystem (`reconai/ingest/`)
- `base_parser.py`: Robust date normalization (`parse_flexible_date`), monetary cleaning (`clean_monetary_amount`), safe in-memory file reader (`read_file_bytes_safely`) to bypass Windows Excel locks, and `get_excel_sheets_info()` to extract all sheet metadata.
- `statement_parser.py`: Multi-page tabular extraction with automatic column detection (Date, Narration, Cheque/Ref, Debit, Credit, Balance), header boundary detection across 100+ pages, bank account number extraction (`Account No 60453403246`), and fallback to Gemini Vision OCR for scanned statements.
- `ledger_parser.py`: Ingestion of Tally Prime and ERP workbooks. Supports multi-sheet auto-detection (prioritizing the exact matching account sheet like `Sheet2` with 3,796 rows over summary sheets like `Sheet1`), column alignment for 'To'/'By' accounting particulars, Dr/Cr semantics, and index preservation.
- `document_parser.py`: Universal parser for Word tables (`.docx`) via `python-docx` and Multimodal Vision OCR via Gemini 2.5 Flash for phone photos (`.jpg`, `.png`, `.webp`) and handwritten cash vouchers.

### Module 2: Multi-Pass Reconciliation Engine (`reconai/reconcile/`)
- `deterministic.py`:
  - Pass 1A: Exact Amount + Exact Reference/Cheque Number/UTR match.
  - Pass 1B: Exact Amount + Same Date (0 days diff). Handles identical multi-entry amounts on the same date by pairing them 1-to-1.
  - Pass 1C: Exact Amount + Date Timing Window (1 to 5 business days for cheque clearance).
- `fuzzy.py`:
  - Pass 2: Semantic Narration Similarity via RapidFuzz. Strips banking noise (NEFT, RTGS, IMPS, UPI, CHQ, PVT, LTD) and scores core party names. Threshold configurable (default 85%).
- `ai_matcher.py`:
  - Pass 3: Gemini AI semantic matching for complex, highly abbreviated party names.
- `matcher.py`:
  - Diagnostic cause assignment for unmatched items: Direct Bank Charges, Direct Government Credits, Unpresented Cheques, Uncredited Deposits.

### Module 3: Forensic Expense & Compliance Audit Engine (`reconai/audit/`)
- `rules_engine.py`:
  - 🔴 HIGH Severity:
    * Duplicate payment vouchers within 3 days.
    * Payments > ₹50,000 without partner authorization.
    * Invalid 15-character GSTIN format (prevents Input Tax Credit disallowance).
    * Single contract payments > ₹30,000 without Section 194C / 194J TDS deduction notes.
  - 🟡 MEDIUM Severity:
    * Round-sum split payments (₹25k, ₹50k, ₹1 Lakh) indicating unvouched cash advances.
    * Vague / sundry narrations lacking audit trail.
    * Payments to vendors not in the Client Master approved whitelist.
  - 🔵 LOW / INFO:
    * Weekend-dated vouchers (Saturday/Sunday cutoff verifications).
- `audit_manager.py`: Coordinates rules and Gemini semantic flagger.

### Module 4: Client Master Subsystem & Database Persistence (`reconai/models/client.py`, `reconai/db/database.py`)
- ClientProfile: Entity Name, Constitution (Proprietorship, Partnership, Pvt Ltd, Ltd, LLP, Trust), PAN, GSTIN, Address, Contact.
- BankAccountProfile: Bank Name, Account Number, IFSC, Branch, Account Type.
- Pre-approved vendor lists and standing auditor instructions.
- SQLite persistence: Complete CRUD operations, session saving/reloading, and audit trail logs.

### Module 5: Interactive Desktop UI (`reconai/ui/`)
- `main_window.py`: Header bar with interactive Client Dropdown (ComboBox) for instant switching, "🏢 Client Master" editor modal, "🔄 Refresh" button, "💾 Save Session", Theme toggle, and 4-step workflow navigation stepper.
- `ingest_view.py`: Side-by-side Bank Statement & Ledger upload cards. Detects multi-sheet Excel workbooks and displays a `📄 Select Account / Sheet:` dropdown with row counts, plus a `🔄 Reset / Clear Files` button.
- `reconcile_view.py`: Real-time KPI stat cards (Matched %, Probable, Unmatched), dual table view with manual match/unmatch override and Undo/Redo (Ctrl+Z / Ctrl+Y).
- `audit_view.py`: Exception table with severity filter and search. Features a fully editable Forensic Risk & Regulatory Analysis panel allowing auditors to edit findings, adjust severity, and customize recommended CA actions with real-time sync.
- `export_view.py`: Real-time BRS calculation table, editable Auditor Remarks and CA Partner Opinion input fields, and export buttons for 7-Tab Excel and signed PDF.

### Module 6: Deliverables Reporting Subsystem (`reconai/report/`)
- `excel_exporter.py`: 7-tab professional openpyxl workbook:
  1. BRS & Summary (with embedded Auditor Observations & CA Opinion)
  2. Matched Transactions
  3. Probable Matches
  4. Unmatched Bank Statement (with diagnostic causes)
  5. Unmatched Client Books (with diagnostic causes)
  6. Forensic Audit Exceptions
  7. System Audit Trail & Session Metadata
- `pdf_exporter.py`: Executive ReportLab PDF audit report with BRS computation, exception register, Auditor Observations, CA Partner Qualification, and official sign-off blocks.

---

## 4. Standalone Packaging Specification
- PyInstaller `.spec` configuration with hidden imports for PyQt6, docx, reportlab, pdfplumber, google.genai, and sqlite3.
- Output binary: `dist/ReconAI.exe`.
"""

for line in master_prompt_text.strip().split("\n"):
    line_s = line.strip()
    if line_s.startswith("## "):
        add_heading_styled(doc_prompt, line_s[3:], 1)
    elif line_s.startswith("### "):
        add_heading_styled(doc_prompt, line_s[4:], 2)
    elif line_s.startswith("- ") or line_s.startswith("* "):
        bp = doc_prompt.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_before = Pt(2)
        bp.paragraph_format.space_after = Pt(2)
        run_bp = bp.add_run(line_s[2:])
        run_bp.font.name = "Segoe UI"
        run_bp.font.size = Pt(10)
    elif line_s:
        p = doc_prompt.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run_p = p.add_run(line_s)
        run_p.font.name = "Segoe UI"
        run_p.font.size = Pt(10.5)

prompt_path1 = "D:/ReconAI/ReconAI_Complete_Master_Prompt.docx"
prompt_path2 = "C:/Users/DELL/OneDrive/Desktop/ReconAI_Complete_Master_Prompt.docx"
doc_prompt.save(prompt_path1)
try:
    doc_prompt.save(prompt_path2)
    print("Saved Master Prompt to Desktop successfully!")
except Exception as e:
    print("Could not save Master Prompt to Desktop:", e)

print("Saved Master Prompt to workspace successfully:", prompt_path1)

# ==============================================================================
# 2. BUILD FILE 2: ReconAI_Complete_Source_Code.docx
# ==============================================================================
print("Generating File 2: ReconAI_Complete_Source_Code.docx...")
doc_code = docx.Document()
set_page_margins(doc_code)

p2 = doc_code.add_paragraph()
r2 = p2.add_run("ReconAI — Complete Application Source Code")
r2.bold = True
r2.font.name = "Segoe UI"
r2.font.size = Pt(22)
r2.font.color.rgb = PRIMARY_COLOR

p2_sub = doc_code.add_paragraph()
r2_sub = p2_sub.add_run("Full Codebase Repository Export (All 47 Python, UI, Database, Parser, Model & Config Files)")
r2_sub.font.name = "Segoe UI"
r2_sub.font.size = Pt(12)
r2_sub.font.italic = True
r2_sub.font.color.rgb = RGBColor(100, 116, 139)

add_callout_box(doc_code, "Codebase Repository Information", "This document contains the complete, unabridged source code for all modules of ReconAI. Each file is presented with its exact file path, description, and complete code content ready for copy-pasting or automated code deployment.")

code_files_list = []
for root, dirs, files in os.walk("reconai"):
    for f in sorted(files):
        if f.endswith(".py"):
            code_files_list.append(os.path.join(root, f))

for root_f in ["run.py", "requirements.txt", "config.json", "packaging/reconai.spec"]:
    if os.path.exists(root_f):
        code_files_list.append(root_f)

print(f"Total source code files to render: {len(code_files_list)}")

for file_idx, file_path in enumerate(code_files_list, 1):
    norm_path = file_path.replace("\\", "/")
    print(f"  [{file_idx}/{len(code_files_list)}] Adding {norm_path}...")
    
    # Heading for each file
    add_heading_styled(doc_code, f"{file_idx}. {norm_path}", 1)
    
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f_in:
            content = f_in.read()
    except Exception as err:
        content = f"# Error reading file: {err}"
        
    line_count = len(content.splitlines())
    
    # Meta info
    p_meta = doc_code.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(0)
    p_meta.paragraph_format.space_after = Pt(4)
    r_meta = p_meta.add_run(f"Path: {norm_path} | Lines of Code: {line_count} | Size: {os.path.getsize(file_path):,} bytes")
    r_meta.font.name = "Segoe UI"
    r_meta.font.size = Pt(9.5)
    r_meta.font.color.rgb = RGBColor(100, 116, 139)

    # Code Table Box
    tbl = doc_code.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(7.0)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{LIGHT_BG_HEX}"/>')
    tcPr.append(shd)
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="12" w:space="0" w:color="94A3B8"/><w:top w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/><w:right w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/><w:bottom w:val="single" w:sz="12" w:space="0" w:color="CBD5E1"/></w:tcBorders>')
    tcPr.append(borders)
    
    p_code = cell.paragraphs[0]
    p_code.paragraph_format.space_before = Pt(4)
    p_code.paragraph_format.space_after = Pt(4)
    r_code = p_code.add_run(content)
    r_code.font.name = "Consolas"
    r_code.font.size = Pt(8.5)
    r_code.font.color.rgb = RGBColor(15, 23, 42)
    
    doc_code.add_paragraph()  # Spacing

code_path1 = "D:/ReconAI/ReconAI_Complete_Source_Code.docx"
code_path2 = "C:/Users/DELL/OneDrive/Desktop/ReconAI_Complete_Source_Code.docx"
doc_code.save(code_path1)
try:
    doc_code.save(code_path2)
    print("Saved Source Code document to Desktop successfully!")
except Exception as e:
    print("Could not save Source Code document to Desktop:", e)

print("Saved Source Code document to workspace successfully:", code_path1)
print("ALL SEPARATE WORD FILES CREATED SUCCESSFULLY!")

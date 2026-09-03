import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

doc = docx.Document()

# Set standard margins (0.75 in)
for s in doc.sections:
    s.top_margin = Inches(0.75)
    s.bottom_margin = Inches(0.75)
    s.left_margin = Inches(0.75)
    s.right_margin = Inches(0.75)

PRIMARY_COLOR = RGBColor(30, 58, 138)     # #1E3A8A Dark Navy
SECONDARY_COLOR = RGBColor(37, 99, 235)   # #2563EB Blue
DARK_GRAY = RGBColor(51, 65, 85)          # #334155
LIGHT_BG_HEX = "F1F5F9"
NAVY_BG_HEX = "1E3A8A"

def style_heading(p, text, level):
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Segoe UI"
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = PRIMARY_COLOR
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = SECONDARY_COLOR
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = DARK_GRAY
    return p

def add_callout(doc, title, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(7.0)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{LIGHT_BG_HEX}"/>')
    tcPr.append(shd)
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="2563EB"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(borders)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r_title = p.add_run(f"📌 {title}\n")
    r_title.bold = True
    r_title.font.name = "Segoe UI"
    r_title.font.color.rgb = PRIMARY_COLOR
    r_text = p.add_run(text)
    r_text.font.name = "Segoe UI"
    r_text.font.size = Pt(10.5)

# Document Header Title
p_title = doc.add_paragraph()
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after = Pt(2)
run_title = p_title.add_run("ReconAI — Bank Reconciliation & Expense Audit Assistant")
run_title.bold = True
run_title.font.name = "Segoe UI"
run_title.font.size = Pt(24)
run_title.font.color.rgb = PRIMARY_COLOR

p_sub = doc.add_paragraph()
p_sub.paragraph_format.space_after = Pt(16)
run_sub = p_sub.add_run("Standard Operating Procedure & Complete CA User Manual (Updated: September 2026)")
run_sub.font.name = "Segoe UI"
run_sub.font.size = Pt(12)
run_sub.font.italic = True
run_sub.font.color.rgb = RGBColor(100, 116, 139)

add_callout(doc, "About ReconAI", "ReconAI is an intelligent desktop application engineered specifically for Chartered Accountants, Audit Firms, and Finance Teams. It automates high-speed multi-page Bank Reconciliation (3,700+ transactions in seconds), forensic expense compliance audit checks (Duplicate vouchers, TDS disallowances, GSTIN validations), and generates 7-Tab audit deliverables in Excel and executive sign-off PDF reports.")

# 1. Launching
style_heading(doc.add_paragraph(), "1. How to Launch & Run ReconAI", 1)
p1 = doc.add_paragraph("ReconAI is packaged as a high-performance standalone 64-bit desktop application. You have two options to run the application:")
p1.style.font.name = "Segoe UI"

bp1 = doc.add_paragraph(style="List Bullet")
r = bp1.add_run("Option A: Standalone Executable (Recommended)\n")
r.bold = True
bp1.add_run("Navigate to ")
r_code = bp1.add_run("D:\\ReconAI\\dist\\")
r_code.bold = True
bp1.add_run(" and double-click ")
r_exe = bp1.add_run("ReconAI.exe")
r_exe.bold = True
bp1.add_run(". You can right-click and choose 'Create shortcut' to place a shortcut directly on your Windows Desktop for quick launch.")

bp2 = doc.add_paragraph(style="List Bullet")
r2 = bp2.add_run("Option B: From Terminal / PowerShell\n")
r2.bold = True
bp2.add_run("Open PowerShell or Command Prompt, navigate to the folder, and run:\n")
r_cmd = bp2.add_run("cd D:\\ReconAI\npython run.py")
r_cmd.bold = True
r_cmd.font.name = "Consolas"

# 2. Client Master
style_heading(doc.add_paragraph(), "2. Client Master Data & Instant Dropdown Switching", 1)
doc.add_paragraph("ReconAI maintains a permanent SQLite database (reconai.db) allowing you to save and switch multiple client profiles effortlessly.")

bp_cm1 = doc.add_paragraph(style="List Bullet")
bp_cm1.add_run("Top Header Client Dropdown: ").bold = True
bp_cm1.add_run("Click the scroll-down menu labeled 'Client:' at the top bar. All saved clients (e.g. SUN HEALTH CARE, SUN EYE CARE, etc.) will appear. Click any client to load their legal profile, bank accounts, approved vendor lists, and standing auditor notes.")

bp_cm2 = doc.add_paragraph(style="List Bullet")
bp_cm2.add_run("🏢 Client Master Button: ").bold = True
bp_cm2.add_run("Click this button in the top bar to open the master manager. You can register new clients, specify legal constitution (Proprietorship, Pvt Ltd, LLP, Trust), enter 10-digit PAN and 15-digit GSTIN, link multiple bank accounts, and configure approved vendor whitelists.")

# 3. Step 1: Ingest & Load
style_heading(doc.add_paragraph(), "3. Step 1: Ingest & Load (Bank Statements & Tally Multi-Sheet Data)", 1)
doc.add_paragraph("Step 1 allows you to load both the Bank Statement and Client Ledger. ReconAI supports PDF, Excel (.xlsx, .xls), CSV, Word (.docx), and Phone Photos / Scans.")

add_callout(doc, "CRITICAL: Handling Multi-Sheet Tally Excel Files (3,000+ Items)", "When a client provides a Tally Excel file that contains multiple bank accounts across different sheets (e.g. Sheet1 with 616 rows, Sheet2 with 3,796 rows):\n1. Click 'Browse File' under 2. Client Books / Ledger.\n2. ReconAI will automatically scan the Excel file. If multiple sheets are present, a dropdown '📄 Select Account / Sheet:' will appear!\n3. ReconAI automatically selects the active bank book with the highest volume (e.g. Sheet2: BOM CA 60453403246 (SUN EYE CARE) Book with 3,796 rows).\n4. If you ever need to reconcile a different account from the same file, simply select it from this dropdown before clicking 'Parse & Normalize Files'!")

doc.add_paragraph("Workflow in Step 1:")
s1_p1 = doc.add_paragraph(style="List Number")
s1_p1.add_run("1. Under Bank Statement: ").bold = True
s1_p1.add_run("Click 'Browse File' and pick your bank statement (e.g. 113-page BOM CA BANK STATEMENTS.pdf).")

s1_p2 = doc.add_paragraph(style="List Number")
s1_p2.add_run("2. Under Client Books / Ledger: ").bold = True
s1_p2.add_run("Click 'Browse File' and select your Tally export (e.g. SUN HEALTH CARE HOSPITAL TALLY DATA.xlsx).")

s1_p3 = doc.add_paragraph(style="List Number")
s1_p3.add_run("3. Verify Sheet Selection: ").bold = True
s1_p3.add_run("Ensure the correct sheet is selected in '📄 Select Account / Sheet:' (defaults automatically to the 3,796-row Sheet2).")

s1_p4 = doc.add_paragraph(style="List Number")
s1_p4.add_run("4. Click 'Parse & Normalize Files': ").bold = True
s1_p4.add_run("The engine parses all 3,779 bank statement transactions and 3,787 ledger vouchers in seconds and renders them in the preview tables below.")

s1_p5 = doc.add_paragraph(style="List Number")
s1_p5.add_run("5. Click '▶ Run Automated Reconciliation & Audit': ").bold = True
s1_p5.add_run("ReconAI executes the multi-pass reconciliation engine and expense forensic audit, advancing directly to Step 2.")

# 4. Step 2: Reconciliation
style_heading(doc.add_paragraph(), "4. Step 2: Multi-Pass Reconciliation Engine", 1)
doc.add_paragraph("ReconAI matches transactions using a 3-tier matching cascade:")

bp_r1 = doc.add_paragraph(style="List Bullet")
bp_r1.add_run("Pass 1 — Deterministic Matching: ").bold = True
bp_r1.add_run("Matches 1-to-1 exact amounts on same dates, cheque numbers, UTR references, and timing differences (1 to 5 business days for cheque clearance). Multiple identical entries on the same date are systematically paired 1-to-1.")

bp_r2 = doc.add_paragraph(style="List Bullet")
bp_r2.add_run("Pass 2 — RapidFuzz Semantic Narration Matching: ").bold = True
bp_r2.add_run("Calculates token-sort and token-set similarity across vendor names, ignoring channel noise (NEFT, RTGS, IMPS, UPI, CHQ).")

bp_r3 = doc.add_paragraph(style="List Bullet")
bp_r3.add_run("Pass 3 — Unmatched Diagnostic Analysis: ").bold = True
bp_r3.add_run("Diagnoses unmatched entries into actionable causes: Direct Bank Debits (bank charges, loan EMIs), Direct Deposits (Govt scheme credits), and Outstanding Unpresented Cheques.")

bp_r4 = doc.add_paragraph(style="List Bullet")
bp_r4.add_run("Manual Override & Undo/Redo: ").bold = True
bp_r4.add_run("Select any unmatched statement and ledger entry and click 'Manual Match Pair' or press Ctrl+Z to undo.")

# 5. Step 3: Expense Audit & Editing Forensic Risk
style_heading(doc.add_paragraph(), "5. Step 3: Expense Audit & Editing Forensic Risk Findings", 1)
doc.add_paragraph("Step 3 analyzes all client vouchers for forensic tax and regulatory risks:")

bp_a1 = doc.add_paragraph(style="List Bullet")
bp_a1.add_run("🔴 High Severity: ").bold = True
bp_a1.add_run("Duplicate voucher payments within 3 days, payments > ₹50,000 without partner authorization, invalid 15-digit GSTIN formats (loss of Input Tax Credit), and contract payments > ₹30,000 lacking Section 194C/194J TDS compliance.")

bp_a2 = doc.add_paragraph(style="List Bullet")
bp_a2.add_run("🟡 Medium Severity: ").bold = True
bp_a2.add_run("Large round-figure sums (₹25k, ₹50k, ₹1 Lakh) indicating unvouched cash advances, vague descriptions ('Office Exp', 'Misc'), and unapproved vendors.")

bp_a3 = doc.add_paragraph(style="List Bullet")
bp_a3.add_run("🔵 Low / Info: ").bold = True
bp_a3.add_run("Weekend transactions (Saturday/Sunday cutoff verifications).")

add_callout(doc, "How to Edit Forensic Risk & Regulatory Findings", "1. Select any row in the exceptions table on the left.\n2. On the right panel, the text boxes are now fully editable!\n3. Edit the '✏️ Auditor Finding / Legal Reasoning' box to add your custom observations or client explanations.\n4. Edit the '✏️ Recommended Action for CA Team' box to customize instructions.\n5. Adjust Severity via the dropdown if needed.\n6. Click '💾 Save & Apply Finding Changes'. Your custom findings will automatically be embedded in the final exported reports!")

# 6. Step 4: Reports & Export
style_heading(doc.add_paragraph(), "6. Step 4: Report & Deliverables Export", 1)
doc.add_paragraph("ReconAI generates two audit deliverables:")

bp_e1 = doc.add_paragraph(style="List Bullet")
bp_e1.add_run("📊 Excel 7-Tab Workbook (.xlsx): ").bold = True
bp_e1.add_run("Features BRS & Summary calculation, Matched Records, Probable Matches, Unmatched Bank, Unmatched Books, Expense Exceptions, and complete Audit Trail.")

bp_e2 = doc.add_paragraph(style="List Bullet")
bp_e2.add_run("📄 Executive CA Sign-Off PDF (.pdf): ").bold = True
bp_e2.add_run("A formal audit memorandum including Executive KPI Summary, Real-Time BRS Statement, Exception Register, Auditor Observations & Professional CA Partner Qualification, with official partner sign-off blocks.")

bp_e3 = doc.add_paragraph(style="List Bullet")
bp_e3.add_run("Auditor Remarks in UI: ").bold = True
bp_e3.add_run("In Step 4, enter your custom observations and partner qualifications before clicking 'Export Excel Report' or 'Export PDF Summary'.")

# 7. Buttons Reference
style_heading(doc.add_paragraph(), "7. UI Buttons & Controls Reference", 1)

table_btns = doc.add_table(rows=1, cols=3)
table_btns.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table_btns.rows[0].cells
hdr_cells[0].text = "Button / Control"
hdr_cells[1].text = "Location"
hdr_cells[2].text = "Function & Action"

for cell in hdr_cells:
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{NAVY_BG_HEX}"/>')
    cell._tc.get_or_add_tcPr().append(shd)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.name = "Segoe UI"

buttons_info = [
    ("🔄 Refresh", "Top Navigation Bar", "Refreshes client master list, syncs active session data, and reloads current views."),
    ("🔄 Reset / Clear Files", "Step 1 (Ingest & Load)", "Clears all uploaded statement & ledger files, resets row counts, and clears preview tables for a fresh run."),
    ("📄 Select Account / Sheet", "Step 1 (Client Ledger)", "Allows switching between multiple bank accounts or sheets in a multi-account Tally Excel file."),
    ("🏢 Client Master", "Top Navigation Bar", "Opens the full Client Master dialog to create, update, or delete client profiles and bank accounts."),
    ("💾 Save Session", "Top Navigation Bar", "Saves the entire active reconciliation session to the SQLite database for future reference."),
    ("💾 Save & Apply Finding Changes", "Step 3 (Expense Audit)", "Saves custom edited auditor reasons and recommended actions into the active audit deliverable."),
]

for b_name, b_loc, b_desc in buttons_info:
    row_cells = table_btns.add_row().cells
    row_cells[0].text = b_name
    row_cells[1].text = b_loc
    row_cells[2].text = b_desc
    for c in row_cells:
        c.paragraphs[0].runs[0].font.name = "Segoe UI"
        c.paragraphs[0].runs[0].font.size = Pt(9.5)

# Save Word Doc in ReconAI folder and on Desktop
out_path1 = "D:/ReconAI/ReconAI_User_Operating_Manual.docx"
out_path2 = "C:/Users/DELL/OneDrive/Desktop/ReconAI_User_Operating_Manual.docx"

doc.save(out_path1)
try:
    doc.save(out_path2)
    print("Saved to Desktop successfully:", out_path2)
except Exception as e:
    print("Could not save to desktop:", e)

print("Saved to workspace successfully:", out_path1)

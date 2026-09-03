import io
import logging
import mimetypes
from pathlib import Path
from typing import Any, Dict, List, Optional, Union
import docx
import pandas as pd

from reconai.models.transaction import StatementTransaction, LedgerEntry, TransactionType
from reconai.ingest.base_parser import parse_flexible_date, clean_monetary_amount
from reconai.ai.client import AIClient

logger = logging.getLogger(__name__)

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tiff"}
DOCX_EXTENSIONS = {".docx", ".doc"}


class DocumentParser:
    """
    Universal document parser handling Word documents (.docx), scanned images
    (JPEG, PNG, WebP), handwritten vouchers, and scanned PDF fallbacks via Gemini Vision OCR.
    """

    def __init__(self, ai_client: Optional[AIClient] = None, api_key: Optional[str] = None):
        self.ai_client = ai_client or AIClient(api_key=api_key)

    def is_image(self, file_path_or_name: Union[str, Path]) -> bool:
        ext = Path(file_path_or_name).suffix.lower()
        return ext in IMAGE_EXTENSIONS

    def is_docx(self, file_path_or_name: Union[str, Path]) -> bool:
        ext = Path(file_path_or_name).suffix.lower()
        return ext in DOCX_EXTENSIONS

    # ================= WORD (.DOCX) INGESTION =================
    def parse_docx_to_dataframe(self, source: Union[str, Path, bytes]) -> pd.DataFrame:
        """Extracts tabular data from Word documents."""
        if isinstance(source, bytes):
            doc = docx.Document(io.BytesIO(source))
        else:
            doc = docx.Document(source)

        all_rows = []
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):
                    all_rows.append(row_data)

        if not all_rows:
            # Fallback: Parse paragraphs line-by-line
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            return pd.DataFrame({"Text": lines})

        # Check maximum column width
        max_cols = max(len(r) for r in all_rows)
        normalized = [r + [""] * (max_cols - len(r)) for r in all_rows]

        # Use first non-empty row as header
        header = normalized[0]
        data = normalized[1:]
        return pd.DataFrame(data, columns=header)

    # ================= MULTIMODAL VISION INGESTION (IMAGES / SCANS) =================
    def extract_statement_from_image(
        self,
        source: Union[str, Path, bytes],
        file_name: Optional[str] = None,
    ) -> List[StatementTransaction]:
        """Extracts bank statement transactions from scanned photo, screenshot, or phone image."""
        if not self.ai_client.is_available:
            raise ValueError("Gemini API key is required to parse scanned images or handwritten copies.")

        file_bytes, mime_type = self._get_bytes_and_mime(source, file_name or "statement.png")

        prompt = (
            "You are an expert financial OCR model. Analyze this bank statement image or scanned document.\n"
            "Extract every single transaction line with 100% accuracy, including date, full description/narration, amount, type (DEBIT or CREDIT), and running balance.\n"
            "Output ONLY a valid JSON list of objects matching this exact schema:\n"
            '[{\n'
            '  "source_row_ref": 1,\n'
            '  "date": "YYYY-MM-DD",\n'
            '  "description": "...",\n'
            '  "amount": "1234.50",\n'
            '  "type": "DEBIT"|"CREDIT",\n'
            '  "balance": "5678.90"\n'
            '}]'
        )

        items = self.ai_client.generate_json_multimodal(
            prompt=prompt,
            mime_type=mime_type,
            file_bytes=file_bytes,
            system_instruction="You are an expert banking document OCR assistant. Extract financial tables from images into structured JSON."
        )

        if not isinstance(items, list):
            return []

        transactions: List[StatementTransaction] = []
        for idx, item in enumerate(items):
            dt = parse_flexible_date(item.get("date"))
            amt = clean_monetary_amount(item.get("amount"))
            bal = clean_monetary_amount(item.get("balance")) if item.get("balance") else None
            t_str = str(item.get("type", "DEBIT")).upper()
            ttype = TransactionType.CREDIT if "CR" in t_str else TransactionType.DEBIT

            if dt and amt:
                transactions.append(
                    StatementTransaction(
                        source_row_ref=item.get("source_row_ref", idx + 1),
                        date=dt,
                        description=str(item.get("description", "Scanned Transaction")),
                        amount=amt,
                        type=ttype,
                        balance=bal,
                        raw_data=item,
                    )
                )

        return transactions

    def extract_ledger_from_image(
        self,
        source: Union[str, Path, bytes],
        file_name: Optional[str] = None,
    ) -> List[LedgerEntry]:
        """Extracts ledger entries from scanned handwritten records, paper books, or voucher photos."""
        if not self.ai_client.is_available:
            raise ValueError("Gemini API key is required to parse scanned images or handwritten copies.")

        file_bytes, mime_type = self._get_bytes_and_mime(source, file_name or "ledger.png")

        prompt = (
            "You are an expert forensic accountant OCR assistant. Analyze this client accounting book, voucher, or handwritten ledger image.\n"
            "Extract all expense payments, receipts, or journal entries.\n"
            "Output ONLY a valid JSON list of objects matching this exact schema:\n"
            '[{\n'
            '  "source_row_ref": "VCH-01",\n'
            '  "date": "YYYY-MM-DD",\n'
            '  "description": "Particulars / Narration",\n'
            '  "amount": "1234.50",\n'
            '  "type": "DEBIT"|"CREDIT",\n'
            '  "account_name": "Party or Account Head",\n'
            '  "voucher_type": "Payment"|"Receipt"|"Journal",\n'
            '  "approver_ref": "Approver name or initials if present",\n'
            '  "gstin": "15-digit GSTIN if visible"\n'
            '}]'
        )

        items = self.ai_client.generate_json_multimodal(
            prompt=prompt,
            mime_type=mime_type,
            file_bytes=file_bytes,
            system_instruction="You are an expert forensic CA document auditor. Extract handwritten and typed accounting ledger entries with high precision."
        )

        if not isinstance(items, list):
            return []

        entries: List[LedgerEntry] = []
        for idx, item in enumerate(items):
            dt = parse_flexible_date(item.get("date"))
            amt = clean_monetary_amount(item.get("amount"))
            t_str = str(item.get("type", "DEBIT")).upper()
            ttype = TransactionType.CREDIT if "CR" in t_str else TransactionType.DEBIT

            if dt and amt:
                entries.append(
                    LedgerEntry(
                        source_row_ref=str(item.get("source_row_ref", idx + 1)),
                        date=dt,
                        description=str(item.get("description", "Scanned Ledger Entry")),
                        amount=amt,
                        type=ttype,
                        account_name=item.get("account_name"),
                        voucher_type=item.get("voucher_type"),
                        approver_ref=item.get("approver_ref"),
                        gstin=item.get("gstin"),
                        raw_data=item,
                    )
                )

        return entries

    def _get_bytes_and_mime(self, source: Union[str, Path, bytes], filename: str) -> tuple[bytes, str]:
        if isinstance(source, bytes):
            data = source
        else:
            with open(source, "rb") as f:
                data = f.read()

        mime_type, _ = mimetypes.guess_type(filename)
        if not mime_type:
            if filename.lower().endswith(".png"):
                mime_type = "image/png"
            elif filename.lower().endswith((".jpg", ".jpeg")):
                mime_type = "image/jpeg"
            elif filename.lower().endswith(".webp"):
                mime_type = "image/webp"
            elif filename.lower().endswith(".pdf"):
                mime_type = "application/pdf"
            else:
                mime_type = "image/png"

        return data, mime_type
